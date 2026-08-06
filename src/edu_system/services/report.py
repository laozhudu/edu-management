"""
报表生成服务 — 从数据库直接生成 Excel/Word 报表
"""

from collections import defaultdict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from sqlalchemy.orm import Session

from edu_system.models import Exam, Semester
from edu_system.services.score import ScoreService

# ── 样式常量 (Excel) ──
THIN = Border(left=Side("thin"), right=Side("thin"), top=Side("thin"), bottom=Side("thin"))
F_TITLE = Font(name="宋体", bold=True, size=14)
F_HEAD = Font(name="宋体", bold=True, size=10)
F_NORM = Font(name="宋体", size=9)
F_SMALL = Font(name="宋体", size=8)
FIL_H = PatternFill("solid", fgColor="D9E1F2")
AC = Alignment(horizontal="center", vertical="center", wrap_text=True)
AL = Alignment(horizontal="left", vertical="center", wrap_text=True)


def _sc(ws, r, c, v, font=None, fill=None, align=None):
    cell = ws.cell(row=r, column=c, value=v)
    if font:
        cell.font = font
    if fill:
        cell.fill = fill
    cell.alignment = align or AC
    cell.border = THIN
    return cell


def _merge(ws, r1, c1, r2, c2, v, font=None, fill=None, align=None):
    ws.merge_cells(start_row=r1, start_column=c1, end_row=r2, end_column=c2)
    return _sc(ws, r1, c1, v, font, fill, align)


class ReportService:
    """报表引擎 — 支持 Excel 和 Word 双格式输出"""

    def __init__(self, session: Session):
        self.session = session
        self.score_svc = ScoreService(session)

    # ═══════════════════════════════════
    # 1. 考试标准报表 (Excel) - 保持原有功能
    # ═══════════════════════════════════

    def generate_exam_report(self, exam_id: int, output_path: str) -> str:
        """生成一次考试的标准报表（原始成绩 + 科分析 + 排名）"""
        students, subjects, configs = self.score_svc.get_exam_scores(exam_id)
        ranked = self.score_svc.calc_grade_ranks(exam_id)

        exam = self.session.get(Exam, exam_id)
        if not exam:
            raise ValueError(f"考试不存在: exam_id={exam_id}")

        semester_label = exam.semester.display_label if exam.semester else ""
        grade_name = exam.grade.name if exam.grade else ""

        classes = sorted(set(s["class_name"] for s in ranked))
        wb = Workbook()

        # ── Sheet 1: 原始成绩 ──
        ws = wb.active
        ws.title = "原始成绩"
        headers = ["班级", "座号", "姓名"] + subjects + ["总分", "平均", "班名", "级名"]
        _merge(
            ws,
            1,
            1,
            1,
            len(headers),
            f"{semester_label}  {grade_name}  {exam.name}  成绩表",
            F_TITLE,
        )
        for c, h in enumerate(headers, 1):
            _sc(ws, 3, c, h, F_HEAD, FIL_H, AC)

        for i, s in enumerate(ranked, 4):
            _sc(ws, i, 1, s["class_name"], F_NORM, align=AC)
            _sc(ws, i, 2, s.get("student_no", ""), F_NORM, align=AC)
            _sc(ws, i, 3, s["name"], F_NORM)
            for si, subj in enumerate(subjects):
                v = s["scores"].get(subj)
                _sc(ws, i, 4 + si, v if v is not None else "", F_NORM, align=AC)
            _sc(ws, i, 4 + len(subjects), s.get("total", ""), F_NORM, align=AC)
            _sc(ws, i, 5 + len(subjects), s.get("grade_rank", ""), F_NORM, align=AC)

        # ── Sheet 2: 科分析 ──
        ws2 = wb.create_sheet("科分析")
        _merge(ws2, 1, 1, 1, 8, f"{semester_label}  {grade_name}  {exam.name}  各科统计", F_TITLE)
        row = 4
        for subj in subjects:
            _merge(ws2, row, 1, row, 8, subj, F_HEAD, FIL_H, AC)
            row += 1
            stat_h = ["班级", "人数", "平均分", "及格%", "良好%", "优秀%", "低分%"]
            for c, h in enumerate(stat_h, 1):
                _sc(ws2, row, c, h, F_HEAD, FIL_H, AC)
            row += 1

            for cls in classes:
                scores = [
                    s["scores"].get(subj)
                    for s in ranked
                    if s["class_name"] == cls and s["scores"].get(subj) is not None
                ]
                if not scores:
                    continue
                cfg = configs.get(subj, {"full_mark": 100, "pass_line": 60})
                pl = cfg.get("pass_line", 60)
                valid = [x for x in scores if x is not None]
                _sc(ws2, row, 1, f"{cls}班", F_NORM, align=AC)
                _sc(ws2, row, 2, len(valid), F_NORM, align=AC)
                _sc(ws2, row, 3, round(sum(valid) / len(valid), 2), F_NORM, align=AC)
                _sc(
                    ws2,
                    row,
                    4,
                    round(sum(1 for x in valid if x >= pl) / len(valid) * 100, 1),
                    F_NORM,
                    align=AC,
                )
                _sc(
                    ws2,
                    row,
                    5,
                    round(sum(1 for x in valid if x >= pl * 1.2) / len(valid) * 100, 1),
                    F_NORM,
                    align=AC,
                )
                _sc(
                    ws2,
                    row,
                    6,
                    round(sum(1 for x in valid if x >= pl * 1.4) / len(valid) * 100, 1),
                    F_NORM,
                    align=AC,
                )
                _sc(
                    ws2,
                    row,
                    7,
                    round(sum(1 for x in valid if x < pl * 0.5) / len(valid) * 100, 1),
                    F_NORM,
                    align=AC,
                )
                row += 1
            row += 1

        # ── Sheet 3: 全级排 ──
        ws3 = wb.create_sheet("全级排")
        _merge(
            ws3,
            1,
            1,
            1,
            4 + len(subjects),
            f"{semester_label}  {grade_name}  {exam.name}  按总分名次排列",
            F_TITLE,
        )
        h3 = ["级名次", "班级", "姓名"] + subjects + ["总分"]
        for c, h in enumerate(h3, 1):
            _sc(ws3, 3, c, h, F_HEAD, FIL_H, AC)
        for i, s in enumerate(ranked, 4):
            _sc(ws3, i, 1, s.get("grade_rank", ""), F_NORM, align=AC)
            _sc(ws3, i, 2, s["class_name"], F_NORM, align=AC)
            _sc(ws3, i, 3, s["name"], F_NORM)
            for si, subj in enumerate(subjects):
                v = s["scores"].get(subj)
                _sc(ws3, i, 4 + si, v if v is not None else "", F_NORM, align=AC)
            _sc(ws3, i, 4 + len(subjects), s.get("total", ""), F_NORM, align=AC)

        wb.save(output_path)
        return output_path

    # ═══════════════════════════════════
    # 2. 学籍变动情况表 (Excel) - 保持原有功能
    # ═══════════════════════════════════

    def generate_change_report(self, semester_id: int, output_path: str) -> str:
        """生成学籍变动情况表（带公文格式）"""

        sem = self.session.get(Semester, semester_id) if semester_id else None
        if not sem:
            raise ValueError("学期不存在")

        wb = Workbook()
        wb.remove(wb.active)
        f14 = Font(name="宋体", size=14)
        f12 = Font(name="宋体", size=12)

        grades_data = [
            ("初一级", "1", "初一(本校）"),
            ("初二级", "2", "初二 (本校）"),
            ("初三级", "3", "初三（本校）"),
        ]

        for grade_name, grade_prefix, sheet_name in grades_data:
            ws = wb.create_sheet(title=sheet_name)
            for ci, w in enumerate([4.5, 6, 4.4, 4.4, 4.4, 4.4, 4.4, 28, 32]):
                ws.column_dimensions[chr(65 + ci)].width = w
            ws.row_dimensions[1].height = 40

            classes = self.session.execute(
                "SELECT c.id, c.name FROM classes c JOIN grades g ON c.grade_id=g.id "
                "WHERE g.name LIKE :gname ORDER BY c.name",
                {"gname": f"%{grade_name}%"},
            ).all()

            # 标题
            ws.merge_cells("A1:I1")
            ws[
                "A1"
            ].value = f"{sem.year_start}-{sem.year_start + 1}学年度{sem.semester}{grade_name}学生变动情况表"
            ws["A1"].font = f14
            ws["A1"].alignment = AC

            # 表头
            row = 3
            for cr, vl in [
                ("A3", "班级"),
                ("B3", "原有人数"),
                ("E3", "现有人数"),
                ("H3", "变动情况"),
            ]:
                ws[cr].value = vl
                ws[cr].font = f12
                ws[cr].alignment = AC

            row = 4
            for cl in classes:
                ws.row_dimensions[row].height = 36
                counts = self.session.execute(
                    "SELECT COUNT(*), SUM(CASE WHEN gender='男' THEN 1 ELSE 0 END) "
                    "FROM students WHERE class_id=? AND status='在校'",
                    (cl.id,),
                ).first()
                ct, cm = (counts[0] or 0), (counts[1] or 0)
                cf = ct - cm

                ws.cell(row=row, column=1, value=f"{cl.name}班")
                ws.cell(row=row, column=2, value=ct)
                ws.cell(row=row, column=3, value=cm)
                ws.cell(row=row, column=4, value=cf)
                ws.cell(row=row, column=5, value=ct)
                row += 1

        wb.save(output_path)
        return output_path

    # ═══════════════════════════════════
    # 3. 成绩单模版批量生成 (Word/Excel) — 新增核心功能
    # ═══════════════════════════════════

    def generate_report_cards_word(
        self,
        exam_id: int,
        output_dir: str,
        template_path: str | None = None,
        single_file: bool = True,
    ) -> list[str]:
        """
        批量生成成绩单 Word 文档（每人一份或合并为单文件）

        Args:
            exam_id: 考试 ID
            output_dir: 输出目录
            template_path: Word 模版路径，None 使用内置模版
            single_file: True=合并为单个 docx，False=每人单独文件

        Returns:
            生成的文件路径列表
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        students, subjects, configs = self.score_svc.get_exam_scores(exam_id)
        ranked = self.score_svc.calc_grade_ranks(exam_id)

        exam = self.session.get(Exam, exam_id)
        if not exam:
            raise ValueError(f"考试不存在: exam_id={exam_id}")

        semester_label = exam.semester.display_label if exam.semester else ""
        grade_name = exam.grade.name if exam.grade else ""

        # 按班级分组
        from collections import defaultdict

        by_class = defaultdict(list)
        for s in ranked:
            by_class[s["class_name"]].append(s)

        generated_files = []

        if single_file:
            # 合并为单个文档
            doc = Document()
            self._setup_doc_style(doc)
            first = True

            for cls_name in sorted(by_class.keys()):
                cls_students = by_class[cls_name]
                if not first:
                    doc.add_page_break()
                first = False

                # 标题
                title = doc.add_paragraph()
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = title.add_run(f"{semester_label} {grade_name} {exam.name} 成绩单")
                run.font.size = Pt(16)
                run.bold = True

                # 班级信息
                info = doc.add_paragraph()
                info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = info.add_run(f"班级：{cls_name}  学期：{semester_label}  年级：{grade_name}")
                run.font.size = Pt(11)

                doc.add_paragraph()  # 空行

                # 成绩表格
                table = doc.add_table(rows=1, cols=len(subjects) + 7, style="Table Grid")
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                self._set_table_font(table, "宋体", Pt(9))

                # 表头
                hdr = table.rows[0]
                headers = ["班级", "座号", "姓名"] + subjects + ["总分", "平均", "班名", "级名"]
                for i, h in enumerate(headers):
                    cell = hdr.cells[i]
                    cell.text = h
                    self._format_cell(cell, bold=True, size=Pt(9), align="center")

                # 数据行
                cls_students = by_class[cls_name]
                for s in cls_students:
                    row = table.add_row()
                    row.cells[0].text = s["class_name"]
                    row.cells[1].text = str(s.get("student_no", ""))
                    row.cells[2].text = s["name"]
                    total = 0
                    for si, subj in enumerate(subjects):
                        v = s["scores"].get(subj)
                        row.cells[3 + si].text = str(v) if v is not None else ""
                        if v is not None:
                            total += v
                    row.cells[3 + len(subjects)].text = str(s.get("total", ""))
                    row.cells[4 + len(subjects)].text = str(s.get("avg", ""))
                    row.cells[5 + len(subjects)].text = str(s.get("class_rank", ""))
                    row.cells[6 + len(subjects)].text = str(s.get("grade_rank", ""))
                    self._format_row(row.cells, Pt(9))

                doc.add_paragraph()  # 空行

            # 保存
            output_path = f"{output_dir}/{exam.name}_成绩单.docx"
            doc.save(output_path)
            generated_files.append(output_path)
        else:
            # 每人单独文件
            for cls_name in sorted(by_class.keys()):
                cls_students = by_class[cls_name]
                for s in cls_students:
                    doc = Document()
                    self._setup_doc_style(doc)

                    # 标题
                    title = doc.add_paragraph()
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = title.add_run(f"{semester_label} {grade_name} {exam.name} 成绩单")
                    run.font.size = Pt(16)
                    run.bold = True

                    # 学生信息
                    info = doc.add_paragraph()
                    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = info.add_run(
                        f"班级：{s['class_name']}  座号：{s.get('student_no', '')}  姓名：{s['name']}"
                    )
                    run.font.size = Pt(11)

                    doc.add_paragraph()

                    # 成绩表
                    table = doc.add_table(rows=1, cols=len(subjects) + 3, style="Table Grid")
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    self._set_table_font(table, "宋体", Pt(9))

                    # 表头
                    hdr = table.rows[0]
                    headers = ["科目", "成绩"] + ["总分", "平均", "班名", "级名"]
                    for i, h in enumerate(headers):
                        cell = hdr.cells[i]
                        cell.text = h
                        self._format_cell(cell, bold=True, size=Pt(9), align="center")

                    # 数据
                    total = 0
                    for subj in subjects:
                        row = table.add_row()
                        v = s["scores"].get(subj)
                        row.cells[0].text = subj
                        row.cells[1].text = str(v) if v is not None else ""
                        if v is not None:
                            total += v
                        self._format_row(row.cells, Pt(9))

                    row = table.add_row()
                    row.cells[0].text = "总分"
                    row.cells[1].text = str(s.get("total", ""))
                    row.cells[2].text = str(s.get("avg", ""))
                    row.cells[3].text = str(s.get("class_rank", ""))
                    row.cells[4].text = str(s.get("grade_rank", ""))

                    # 保存
                    safe_name = f"{s['class_name']}_{s['name']}".replace("/", "_")
                    output_path = f"{output_dir}/{exam.name}_{safe_name}_成绩单.docx"
                    doc.save(output_path)
                    generated_files.append(output_path)

        return generated_files

    def _get_subject_cols(self, si):
        """获取列数"""
        return len(subjects) + 4  # 班级、座号、姓名、10科目、总分、平均、班名、级名

    def _setup_doc_style(self, doc: Document):
        """设置文档默认样式"""
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10)

    def _set_table_font(self, table, font_name, size):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
                        run.font.size = size

    def _format_cell(self, cell, bold=False, size=Pt(9), align="center"):
        for paragraph in cell.paragraphs:
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
            )
            for run in paragraph.runs:
                run.bold = bold
                run.font.size = size
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _format_row(self, cells, size):
        for cell in cells:
            self._format_cell(cell, size=size)

    # 兼容旧接口
    def generate_report_cards(
        self,
        exam_id: int,
        output_dir: str,
        template_path: str | None = None,
        single_file: bool = True,
    ) -> list[str]:
        """生成成绩单 - 兼容旧接口"""
        return self.generate_report_cards_word(
            exam_id=exam_id,
            output_dir=output_dir,
            template_path=template_path,
            single_file=single_file,
        )

    # ═══════════════════════════════════
    # 4. 成绩单 Excel 模版导出
    # ═══════════════════════════════════

    def generate_report_cards_excel(
        self,
        exam_id: int,
        output_dir: str,
        single_file: bool = True,
    ) -> list[str]:
        """
        批量导出成绩单 Excel（每人一行或每人一表）
        """
        students, subjects, configs = self.score_svc.get_exam_scores(exam_id)
        ranked = self.score_svc.calc_grade_ranks(exam_id)

        exam = self.session.get(Exam, exam_id)
        if not exam:
            raise ValueError(f"考试不存在: exam_id={exam_id}")

        semester_label = exam.semester.display_label if exam.semester else ""
        grade_name = exam.grade.name if exam.grade else ""

        by_class = defaultdict(list)
        for s in ranked:
            by_class[s["class_name"]].append(s)

        generated = []

        if single_file:
            wb = Workbook()
            wb.remove(wb.active)

            for cls_name in sorted(by_class.keys()):
                cls_students = by_class[cls_name]
                ws = wb.create_sheet(title=cls_name[:31])

                # 标题行
                ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(subjects) + 4)
                _sc(
                    ws,
                    1,
                    1,
                    f"{semester_label} {grade_name} {exam.name} {cls_name}班 成绩单",
                    F_TITLE,
                    align=AC,
                )

                # 表头
                headers = ["班级", "座号", "姓名"] + subjects + ["总分", "平均", "班名", "级名"]
                for c, h in enumerate(headers, 1):
                    _sc(ws, 3, c, h, F_HEAD, FIL_H, AC)

                for i, s in enumerate(cls_students, 4):
                    _sc(ws, i, 1, s["class_name"], F_NORM, align=AC)
                    _sc(ws, i, 2, s.get("student_no", ""), F_NORM, align=AC)
                    _sc(ws, i, 3, s["name"], F_NORM)
                    total = 0
                    for si, subj in enumerate(subjects):
                        v = s["scores"].get(subj)
                        _sc(ws, i, 4 + si, v if v is not None else "", F_NORM, align=AC)
                        if v is not None:
                            total += v
                    _sc(ws, i, 4 + len(subjects), s.get("total", ""), F_NORM, align=AC)
                    _sc(ws, i, 5 + len(subjects), s.get("avg", ""), F_NORM, align=AC)
                    _sc(ws, i, 5 + len(subjects), s.get("class_rank", ""), F_NORM, align=AC)
                    _sc(ws, i, 6 + len(subjects), s.get("grade_rank", ""), F_NORM, align=AC)

            output_path = f"{output_dir}/{exam.name}_成绩单.xlsx"
            wb.save(output_path)
            generated.append(output_path)
        else:
            # 每人一个文件
            for s in ranked:
                wb = Workbook()
                ws = wb.active
                ws.title = "成绩单"

                _merge(
                    ws,
                    1,
                    1,
                    1,
                    len(self._get_subject_cols(si)) + 3,
                    f"{semester_label} {grade_name} {exam.name} 成绩单",
                    F_TITLE,
                    align=AC,
                )

                # 学生信息
                _sc(ws, 3, 1, "班级:", F_NORM)
                _sc(ws, 3, 2, s["class_name"], F_NORM)
                _sc(ws, 3, 3, "座号:", F_NORM)
                _sc(ws, 3, 4, s.get("student_no", ""), F_NORM)
                _sc(ws, 3, 5, "姓名:", F_NORM)
                _sc(ws, 3, 6, s["name"], F_NORM)

                headers = ["科目", "成绩"] + ["总分", "平均", "班名", "级名"]
                for c, h in enumerate(headers, 1):
                    _sc(ws, 5, c, h, F_HEAD, FIL_H, AC)

                total = 0
                for si, subj in enumerate(subjects):
                    row = 6 + si
                    v = s["scores"].get(subj)
                    _sc(ws, row, 1, subj, F_NORM)
                    _sc(ws, row, 2, v if v is not None else "", F_NORM, align=AC)
                    if v is not None:
                        total += v

                _sc(ws, 6 + len(subjects), 1, "总分", F_NORM)
                _sc(ws, 6 + len(subjects), 2, s.get("total", ""), F_NORM)
                _sc(ws, 6 + len(subjects), 3, s.get("avg", ""), F_NORM)
                _sc(ws, 6 + len(subjects), 4, s.get("class_rank", ""), F_NORM)
                _sc(ws, 6 + len(subjects), 5, s.get("grade_rank", ""), F_NORM)

                output_path = f"{output_dir}/{exam.name}_{s['class_name']}_{s['name']}_成绩单.xlsx"
                wb.save(output_path)
                generated.append(output_path)

        return generated

    # 兼容旧接口
    def generate_report_cards(
        self,
        exam_id: int,
        output_dir: str,
        template_path: str | None = None,
        single_file: bool = True,
    ) -> list[str]:
        """生成成绩单 - 兼容旧接口，默认生成 Word"""
        return self.generate_report_cards_word(
            exam_id=exam_id,
            output_dir=output_dir,
            template_path=template_path,
            single_file=single_file,
        )


# ═══════════════════════════════════
# 5. 证书/奖状生成 (Word)
# ═══════════════════════════════════

    def generate_certificate(
        self,
        exam_id: int,
        output_dir: str,
        certificate_type: str = "award",  # "award" 奖状 | "certificate" 证书
        template_path: str | None = None,
        single_file: bool = True,
    ) -> list[str]:
        """
        批量生成证书/奖状 Word 文档

        Args:
            exam_id: 考试 ID
            output_dir: 输出目录
            certificate_type: "award" 奖状 | "certificate" 证书
            template_path: Word 模版路径，None 使用内置模版
            single_file: True=合并为单个 docx，False=每人单独文件

        Returns:
            生成的文件路径列表
        """
        students, subjects, configs = self.score_svc.get_exam_scores(exam_id)
        ranked = self.score_svc.calc_grade_ranks(exam_id)

        exam = self.session.get(Exam, exam_id)
        if not exam:
            raise ValueError(f"考试不存在: exam_id={exam_id}")

        semester_label = exam.semester.display_label if exam.semester else ""
        grade_name = exam.grade.name if exam.grade else ""

        # 筛选获奖学生
        awardees = self._filter_awardees(ranked, certificate_type)
        if not awardees:
            return []

        generated_files = []

        if single_file:
            doc = Document()
            self._setup_doc_style(doc)
            for i, s in enumerate(awardees):
                if i > 0:
                    doc.add_page_break()
                self._add_certificate_page(
                    doc, s, certificate_type, semester_label, grade_name, exam.name
                )
            output_path = f"{output_dir}/{exam.name}_{certificate_type}.docx"
            doc.save(output_path)
            generated_files.append(output_path)
        else:
            for s in awardees:
                doc = Document()
                self._setup_doc_style(doc)
                self._add_certificate_page(
                    doc, s, certificate_type, semester_label, grade_name, exam.name
                )
                safe_name = f"{s['class_name']}_{s['name']}".replace("/", "_")
                output_path = f"{output_dir}/{exam.name}_{certificate_type}_{safe_name}.docx"
                doc.save(output_path)
                generated_files.append(output_path)

        return generated_files

    def _filter_awardees(self, ranked: list[dict], certificate_type: str) -> list[dict]:
        """筛选获奖学生：奖状=班级前30%，证书=总分前20%"""
        if certificate_type == "award":
            by_class = defaultdict(list)
            for s in ranked:
                by_class[s["class_name"]].append(s)
            awardees = []
            for cls_students in by_class.values():
                cutoff = max(1, len(cls_students) * 30 // 100)
                awardees.extend(cls_students[:cutoff])
            return awardees
        else:
            cutoff = max(1, len(ranked) * 20 // 100)
            return ranked[:cutoff]

    def _add_certificate_page(
        self,
        doc: Document,
        student: dict,
        certificate_type: str,
        semester_label: str,
        grade_name: str,
        exam_name: str,
    ):
        """添加一页证书/奖状"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        title_text = "奖 状" if certificate_type == "award" else "结业证书"
        cert_no = (
            f"第 {student['grade_rank']:04d} 号" if certificate_type == "certificate" else ""
        )

        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(title_text)
        run.font.size = Pt(28)
        run.bold = True
        run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)  # 深红

        if cert_no:
            cert_p = doc.add_paragraph()
            cert_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = cert_p.add_run(cert_no)
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

        doc.add_paragraph()

        # 正文
        body = doc.add_paragraph()
        body.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = body.add_run(f"兹证明 {student['name']} 同学")
        run.font.size = Pt(18)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        if certificate_type == "award":
            text = (
                f"在 {exam_name} 中表现优异，位列年级第 {student['grade_rank']} 名，"
                f"班级第 {student['class_rank']} 名，特授予此奖状以资鼓励。"
            )
        else:
            text = (
                f"在 {semester_label} {grade_name} {exam_name} 中成绩优异，"
                f"顺利完成学业，特颁发此证书。"
            )
        body.add_run(text).font.size = Pt(16)
        for run in body.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        doc.add_paragraph()
        doc.add_paragraph()

        # 落款
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(f"{semester_label} {grade_name} {exam_name}\n颁发单位：某某学校")
        run.font.size = Pt(14)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 签名线
        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = sig.add_run("__________________\n校长签名（盖章）")
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 日期
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = date_p.add_run("二〇二四年十二月")
        run.font.size = Pt(14)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


