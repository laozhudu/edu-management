"""
CertificateGenerator 测试（Sprint 4.8.2）
覆盖：docx 渲染（变量/循环）、批量渲染、HTML→PDF、模板缺失异常
"""

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent.parent / "src"))

from docxtpl import DocxTemplate

from edu_system.services.report_certificate import (
    CertificateError,
    CertificateGenerator,
)


@pytest.fixture
def docx_tpl(tmp_path):
    """含变量 + 循环的 docx 模板"""
    tpl_path = tmp_path / "cert.docx"
    tpl = DocxTemplate(str(tpl_path))
    # 先建一个含占位符的模板（docxtpl 用富文本段落）
    from docx import Document

    doc = Document()
    doc.add_paragraph("荣誉证书")
    doc.add_paragraph("兹证明 {{name}} 同学")
    doc.add_paragraph("{% for item in items %}{{ item }}{% endfor %}")
    doc.save(str(tpl_path))
    return tpl_path


@pytest.fixture
def html_tpl(tmp_path):
    tpl_path = tmp_path / "cert.html"
    tpl_path.write_text(
        "<html><body><h1>荣誉证书</h1><p>兹证明 {{ name }} 同学</p></body></html>",
        encoding="utf-8",
    )
    return tpl_path


class TestDocxRender:
    def test_render_docx(self, docx_tpl, tmp_path):
        gen = CertificateGenerator(docx_tpl)
        out = gen.render_docx({"name": "张三", "items": []}, tmp_path / "out.docx")
        assert out.exists()
        # 重新读取验证内容
        doc = DocxTemplate(str(out))
        from docx import Document

        d = Document(str(out))
        texts = [p.text for p in d.paragraphs]
        assert any("张三" in t for t in texts)

    def test_template_not_found(self, tmp_path):
        with pytest.raises(CertificateError, match="不存在"):
            CertificateGenerator(tmp_path / "nonexistent.docx")


class TestHtmlPdf:
    def test_render_pdf(self, html_tpl, tmp_path):
        gen = CertificateGenerator(html_tpl)
        out = gen.render_pdf({"name": "李四"}, tmp_path / "cert.pdf")
        assert out.exists()
        assert out.suffix == ".pdf"
        assert out.stat().st_size > 0

    def test_docx_to_pdf_raises(self, docx_tpl, tmp_path):
        """docx 模板直接转 PDF 提示需 soffice"""
        gen = CertificateGenerator(docx_tpl)
        with pytest.raises(CertificateError, match="soffice"):
            gen.render_pdf({"name": "张三", "items": []}, tmp_path / "out.pdf")


class TestBatch:
    def test_render_batch(self, docx_tpl, tmp_path):
        gen = CertificateGenerator(docx_tpl)
        rows = [{"name": "张三", "items": []}, {"name": "李四", "items": []}]
        outs = gen.render_batch(rows, tmp_path / "out", prefix="cert")
        assert len(outs) == 2
        assert all(p.exists() for p in outs)
        assert outs[0].name == "cert_001.docx"
